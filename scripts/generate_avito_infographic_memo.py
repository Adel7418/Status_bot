import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception as exc:  # pragma: no cover
    raise SystemExit(
        "Не найдена библиотека python-docx. Установите: pip install python-docx"
    ) from exc


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(item, style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(item, style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def build_document() -> Document:
    document = Document()

    # Базовый стиль шрифта
    style = document.styles["Normal"].font
    style.name = "Calibri"
    style.size = Pt(11)

    # Заголовок
    add_heading(
        document,
        "Памятка по составлению описания, фото и инфографики для Авито (ремонт бытовой техники)",
        level=1,
    )
    add_paragraph(document, f"Версия от {datetime.now().strftime('%Y-%m-%d')}")

    # Цели
    add_heading(document, "Цели", level=2)
    add_bullets(
        document,
        [
            "Быстрое считывание ключевого УТП на маленьких превью (1–2 секунды)",
            "Снижение тревоги клиента: гарантия, чек, прозрачность цен",
            "Рост конверсии в обращение и ремонт за счёт релевантности и доверия",
        ],
    )

    # Инфографика — структура
    add_heading(document, "Инфографика — базовая структура", level=2)
    add_bullets(
        document,
        [
            "Крупный УТП: услуга + скорость/доступность (например: ‘Ремонт стиралок на дому • 30–60 мин’)",
            "3–4 буллета преимуществ: ‘Диагностика 0 ₽’, ‘Гарантия до 12 мес’, ‘Оригинальные запчасти’, ‘Оплата по факту’",
            "Якорь цены: ‘от … ₽’ (без перегруза, лучше на второй строке/плашке)",
            "Гео/время: район/город, ‘24/7’, ‘Сегодня’",
            "Визуально: контрастные цвета, крупные цифры, простые иконки, шрифт без засечек",
        ],
    )

    # Уникальность без риска
    add_heading(document, "Как добавить уникальность без риска", level=2)
    add_bullets(
        document,
        [
            "Персонализация: имя мастера, стаж (‘10 лет’), рейтинг (‘4.9★ / 320 отзывов’), фото мастера в форме",
            "Формализованный сервис: чек 54‑ФЗ, договор, гарантийный талон (покажите во 2‑м фото)",
            "Специализация по ‘боли’: NoFrost, инвертор, Direct Drive, Aquastop, типовые симптомы",
            "Прозрачный процесс из 3 шагов: ‘Заявка → Диагностика 0 ₽ → Фикс‑смета до работ’",
            "Цифровые доказательства: ‘>1 200 ремонтов за 12 мес’, ‘Средний выезд 35 мин по [району]’",
            "Прайс-миниатюра вторым фото: 5–7 частых работ ‘от … ₽’ (не грузить главное изображение)",
            "Сезон/контекст: летом — холодильники (утечки), перед праздниками — стиралки (насос)",
            "Честные рамки: ‘Не берёмся за редкий модуль X/Y — скажем сразу’ — повышает доверие",
        ],
    )

    # Описание объявления
    add_heading(document, "Описание объявления — что важнее картинки", level=2)
    add_paragraph(
        document,
        "Первые 120 символов — критичны: услуга + 1–2 триггера (скорость/гарантия/диагностика 0 ₽)",
    )
    add_paragraph(document, "Рекомендуемая структура:", bold=True)
    add_numbered(
        document,
        [
            "Проблема клиента (симптом)",
            "Что делаем (услуга/узлы)",
            "Доказательства (опыт, отзывы, цифры)",
            "Процесс и сроки (приезд, время ремонта)",
            "Прайс ‘от … ₽’",
            "Гарантии/чек/договор",
            "Гео/график работы",
            "Мини‑FAQ (3–4 ответа)",
        ],
    )
    add_paragraph(document, "Мини‑FAQ — снимите ключевые возражения:", bold=True)
    add_bullets(
        document,
        [
            "Сколько стоит выезд/диагностика?",
            "Когда платёж? Что, если не починили?",
            "С какими брендами/узлами работаем?",
            "Что подготовить до приезда (безопасно)?",
        ],
    )

    # Примеры первых строк
    add_heading(document, "Примеры первых строк", level=2)
    add_bullets(
        document,
        [
            "‘Ремонт стиральных машин на дому за 1 визит. Диагностика 0 ₽ при ремонте, чек и гарантия до 12 мес.’",
            "‘Холодильник не морозит? Приезд за 30–60 мин в [район]. Фикс‑смета до работ, оригинальные запчасти.’",
        ],
    )

    # Чек-лист изображения
    add_heading(document, "Чек‑лист изображения", level=2)
    add_bullets(
        document,
        [
            "Один главный посыл + 3–4 буллета",
            "Крупные цифры: ‘30 мин’, ‘0 ₽’, ‘12 мес’, ‘24/7’",
            "Контраст 2–3 цветов, простые иконки",
            "Проверка читабельности в 200–250 px (скрин)",
            "Без перегруза: 8–12 слов на плашке",
            "Без нарушений правил/модерации (контакты на картинке — по правилам площадки)",
        ],
    )

    # Продвижение
    add_heading(document, "Продвижение и поведение объявления", level=2)
    add_bullets(
        document,
        [
            "Вариативность карточек: по симптомам/брендам/узлам",
            "Гео‑укоренение: микрорайоны/станции",
            "Скорость ответа и онлайн‑статус",
            "Инструменты площадки: VIP/поднятия короткими циклами 2–3 дня",
            "Отзывы: просите упоминать ‘симптом + район’",
        ],
    )

    # A/B тест-план
    add_heading(document, "Быстрый A/B тест‑план (1 неделя)", level=2)
    add_numbered(
        document,
        [
            "3 варианта: (1) базовый шаблон, (2) персонализация (фото мастера+рейтинг), (3) симптом‑ориентированный",
            "Метрики: CTR, стоимость обращения, конверсия в выезд/ремонт, средний чек",
            "Равная ротация продвижения; победителя закрепить и доработать описание по возражениям",
        ],
    )

    # Шаблон для инфографики (текст)
    add_heading(document, "Шаблон текста для инфографики", level=2)
    add_bullets(
        document,
        [
            "УТП: ‘Ремонт [техника] на дому • Выезд 30–60 мин’",
            "Буллеты: ‘Диагностика 0 ₽’ • ‘Гарантия до 12 мес’ • ‘Оригинальные запчасти’ • ‘Оплата по факту’",
            "Гео/время: ‘[район/город] • 24/7’",
            "Цена: ‘от … ₽’ (по частым работам во 2‑м фото)",
        ],
    )

    return document


def main() -> None:
    output_dir = os.path.join("docs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Памятка_Авито_ремонт_техники.docx")

    document = build_document()
    document.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()


